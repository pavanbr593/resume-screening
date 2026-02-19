import re
import io


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return clean_text(text)
    except Exception as e:
        return f"[PDF parse error: {e}]"


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return clean_text("\n".join(paragraphs))
    except Exception as e:
        return f"[DOCX parse error: {e}]"


def clean_text(text: str) -> str:
    """Normalize whitespace and remove non-printable characters."""
    # Remove non-printable chars except common whitespace
    text = re.sub(r"[^\x20-\x7E\n\t]", " ", text)
    # Collapse multiple spaces
    text = re.sub(r"[ \t]+", " ", text)
    # Collapse multiple newlines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_resume(filename: str, file_bytes: bytes) -> str:
    """Route to appropriate parser based on file extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif lower.endswith(".docx"):
        return parse_docx(file_bytes)
    else:
        return "[Unsupported file format]"


def compute_resume_quality_score(text: str) -> float:
    """
    Heuristic quality score (0-100) based on:
    - Length
    - Presence of section headers
    - Contact info indicators
    - Presence of quantified achievements
    """
    score = 0.0

    word_count = len(text.split())
    if word_count >= 300:
        score += 25
    elif word_count >= 150:
        score += 15
    else:
        score += 5

    section_keywords = [
        "experience", "education", "skills", "projects",
        "summary", "objective", "certifications", "achievements"
    ]
    text_lower = text.lower()
    found_sections = sum(1 for kw in section_keywords if kw in text_lower)
    score += min(found_sections * 5, 25)

    # Contact info
    has_email = bool(re.search(r"[\w.-]+@[\w.-]+\.\w+", text))
    has_phone = bool(re.search(r"(\+?\d[\d\s\-().]{7,}\d)", text))
    if has_email:
        score += 10
    if has_phone:
        score += 10

    # Quantified achievements (numbers with %)
    quantified = len(re.findall(r"\d+\s*%|\$\s*\d+|\d+\s*(million|billion|k\b)", text_lower))
    score += min(quantified * 5, 30)

    return min(score, 100.0)
